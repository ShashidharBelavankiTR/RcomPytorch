"""
Chunk — Read training files, split into chunks, deduplicate, validate, save JSONL
==================================================================================
Supported file types:
  .txt   — plain text
  .pdf   — text-based PDFs (via pypdf)
  .docx  — Word documents, including tables (via python-docx)
  .xlsx  — Excel spreadsheets, all sheets (via openpyxl)
  .xls   — Legacy Excel format (via openpyxl)
  .md    — Markdown files (syntax stripped before chunking)

Pipeline:
  1. Scan /train recursively for all supported file types
  2. Extract text per file type
  3. Split into overlapping word-count chunks
  4. Deduplicate near-identical chunks (Jaccard similarity on word trigrams)
  5. Validate chunk quality and print a length distribution report
  6. Format each chunk with the model's chat template
  7. Save as JSONL to /data/train.jsonl
"""

import os
import sys
import re
import json
import glob
import hashlib

import config
from model_registry import get_chat_template, format_training_example


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTORS
# ═══════════════════════════════════════════════════════════════════════════════


def extract_text_from_pdf(filepath: str) -> str:
    """Extract all text from a text-based PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        print(f"    [ERROR] pypdf not installed. Fix: pip install pypdf>=4.0.0")
        return ""
    try:
        reader = PdfReader(filepath)
        parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    parts.append(text)
            except Exception as e:
                print(f"    [WARNING] Page {page_num + 1} unreadable in {filepath}: {e}")
        return "\n".join(parts)
    except Exception as e:
        print(f"    [ERROR] Failed to read PDF {filepath}: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    """
    Extract text from a Word .docx file using python-docx.
    Captures paragraph text and table cell contents.
    """
    try:
        from docx import Document
    except ImportError:
        print(f"    [ERROR] python-docx not installed. Fix: pip install python-docx>=1.1.0")
        return ""
    try:
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"    [ERROR] Failed to read DOCX {filepath}: {e}")
        return ""


def extract_text_from_excel(filepath: str) -> str:
    """
    Extract text from an Excel file (.xlsx or .xls) using openpyxl.
    Reads every sheet; each row becomes a pipe-separated text line.
    """
    try:
        import openpyxl
    except ImportError:
        print(f"    [ERROR] openpyxl not installed. Fix: pip install openpyxl>=3.1.0")
        return ""
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[Sheet: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        print(f"    [ERROR] Failed to read Excel {filepath}: {e}")
        return ""


def extract_text_from_markdown(filepath: str) -> str:
    """
    Read a Markdown file and strip common syntax,
    leaving only readable prose for chunking.
    """
    text = extract_text_from_txt(filepath)
    if not text:
        return ""
    text = re.sub(r"```[\s\S]*?```", "", text)           # fenced code blocks
    text = re.sub(r"`[^`\n]+`", "", text)                # inline code
    text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)  # images → alt text
    text = re.sub(r"\[([^\]]*)\]\([^\)]*\)", r"\1", text)    # links → link text
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)  # headers
    text = re.sub(r"\*{1,2}([^*\n]+)\*{1,2}", r"\1", text)  # bold/italic
    text = re.sub(r"_{1,2}([^_\n]+)_{1,2}", r"\1", text)    # underscore bold/italic
    text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)   # blockquotes
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)  # hr
    text = re.sub(r"^[\-\*\+]\s+", "", text, flags=re.MULTILINE)    # bullet lists
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)       # numbered lists
    text = re.sub(r"\n{3,}", "\n\n", text)                          # excess blank lines
    return text.strip()


def get_extractor(filepath: str):
    """Return the correct extractor function based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    return {
        ".txt":      extract_text_from_txt,
        ".pdf":      extract_text_from_pdf,
        ".docx":     extract_text_from_docx,
        ".xlsx":     extract_text_from_excel,
        ".xls":      extract_text_from_excel,
        ".md":       extract_text_from_markdown,
        ".markdown": extract_text_from_markdown,
    }.get(ext)


def extract_text_from_txt(filepath: str) -> str:
    """Read a plain text file, trying common encodings."""
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(filepath, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    print(f"    [WARNING] Could not decode {filepath} with any known encoding — skipped")
    return ""


# ═══════════════════════════════════════════════════════════════════════════════
# CHUNKING
# ═══════════════════════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """
    Split text into overlapping chunks by word count.
    chunk_size:    target words per chunk
    chunk_overlap: words shared between adjacent chunks
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    step = max(1, chunk_size - chunk_overlap)
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunk = " ".join(chunk_words)
        if chunk.strip():
            chunks.append(chunk)
        if i + chunk_size >= len(words):
            break
        i += step
    return chunks


# ═══════════════════════════════════════════════════════════════════════════════
# DEDUPLICATION
# ═══════════════════════════════════════════════════════════════════════════════

def _get_trigrams(text: str) -> frozenset:
    """Return a frozenset of word trigrams (3-word shingles) from text."""
    words = text.lower().split()
    if len(words) < 3:
        return frozenset([" ".join(words)])
    return frozenset(" ".join(words[i:i + 3]) for i in range(len(words) - 2))


def _jaccard(set_a: frozenset, set_b: frozenset) -> float:
    """Compute Jaccard similarity between two sets."""
    if not set_a or not set_b:
        return 0.0
    inter = len(set_a & set_b)
    union = len(set_a | set_b)
    return inter / union if union else 0.0


def deduplicate_chunks(chunks: list[str], threshold: float = 0.85) -> tuple[list[str], int]:
    """
    Remove near-duplicate chunks using two passes:
      Pass 1 — exact deduplication via MD5 hash (O(n), instant)
      Pass 2 — near-duplicate removal via Jaccard similarity on word trigrams

    A length-ratio pre-filter skips pairs that cannot possibly exceed the
    threshold, keeping runtime acceptable for typical dataset sizes.

    Returns: (deduplicated_list, number_of_chunks_removed)
    """
    if not chunks:
        return chunks, 0

    # Pass 1: exact dedup
    seen_hashes: set = set()
    unique: list[str] = []
    exact_removed = 0
    for chunk in chunks:
        h = hashlib.md5(chunk.encode("utf-8", errors="replace")).hexdigest()
        if h in seen_hashes:
            exact_removed += 1
        else:
            seen_hashes.add(h)
            unique.append(chunk)

    # Pass 2: near-dedup
    trigrams = [_get_trigrams(c) for c in unique]
    keep = [True] * len(unique)
    near_removed = 0

    for i in range(len(unique)):
        if not keep[i]:
            continue
        for j in range(i + 1, len(unique)):
            if not keep[j]:
                continue
            # Length-ratio pre-filter: upper bound on Jaccard = min/max
            len_i, len_j = len(trigrams[i]), len(trigrams[j])
            if len_i == 0 or len_j == 0:
                continue
            if min(len_i, len_j) / max(len_i, len_j) < threshold:
                continue
            if _jaccard(trigrams[i], trigrams[j]) >= threshold:
                keep[j] = False
                near_removed += 1

    result = [c for c, k in zip(unique, keep) if k]
    return result, exact_removed + near_removed


# ═══════════════════════════════════════════════════════════════════════════════
# VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════

def validate_chunks(chunks: list[str]) -> dict:
    """
    Analyse chunk quality.
    Returns a report dict with stats and flagged chunks.
    """
    max_words = config.MAX_CHUNK_WORDS or int(config.CHUNK_SIZE * 1.5)
    min_words = config.MIN_CHUNK_WORDS
    word_counts = [len(c.split()) for c in chunks]
    total = len(word_counts)
    if total == 0:
        return {"total": 0}

    sorted_wc = sorted(word_counts)
    mean_wc   = sum(word_counts) / total
    median_wc = sorted_wc[total // 2]
    too_short = [(i, wc) for i, wc in enumerate(word_counts) if wc < min_words]
    too_long  = [(i, wc) for i, wc in enumerate(word_counts) if wc > max_words]

    bucket_edges  = [0, 50, 100, 200, 400, 600, 800, 1024, 1500, 2048, 999_999]
    bucket_labels = [
        "<50", "50–99", "100–199", "200–399", "400–599",
        "600–799", "800–1023", "1024–1499", "1500–2047", "≥2048",
    ]
    buckets = [0] * len(bucket_labels)
    for wc in word_counts:
        for idx in range(len(bucket_edges) - 1):
            if bucket_edges[idx] <= wc < bucket_edges[idx + 1]:
                buckets[idx] += 1
                break

    return {
        "total":     total,
        "mean":      mean_wc,
        "median":    median_wc,
        "min":       sorted_wc[0],
        "max":       sorted_wc[-1],
        "too_short": too_short,
        "too_long":  too_long,
        "buckets":   list(zip(bucket_labels, buckets)),
        "min_words": min_words,
        "max_words": max_words,
    }


def print_validation_report(report: dict):
    """Pretty-print the validation report."""
    if report.get("total", 0) == 0:
        print("  [WARNING] No chunks to validate.")
        return

    total     = report["total"]
    too_short = report["too_short"]
    too_long  = report["too_long"]

    print()
    print("  ── Chunk Quality Report ──────────────────────────────────────")
    print(f"  Total chunks:     {total}")
    print(f"  Mean words:       {report['mean']:.0f}")
    print(f"  Median words:     {report['median']}")
    print(f"  Min words:        {report['min']}")
    print(f"  Max words:        {report['max']}")
    print()
    print("  Word count distribution:")
    for label, count in report["buckets"]:
        if count == 0:
            continue
        bar = "█" * int(30 * count / total)
        print(f"    {label:>10}  {bar:<30}  {count:>5} ({100 * count / total:.1f}%)")
    print()
    if too_short:
        print(f"  [WARN] {len(too_short)} chunk(s) below {report['min_words']} words "
              f"(low information — consider a better source file or lower CHUNK_SIZE)")
    else:
        print(f"  [OK]   No chunks below {report['min_words']} words")
    if too_long:
        print(f"  [WARN] {len(too_long)} chunk(s) above {report['max_words']} words "
              f"(may exceed model context — consider reducing CHUNK_SIZE)")
    else:
        print(f"  [OK]   No chunks above {report['max_words']} words")
    print("  ─────────────────────────────────────────────────────────────")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".xlsx", ".xls", ".md", ".markdown"}


def run_chunking() -> str:
    """
    Full pipeline: scan → extract → chunk → dedup → validate → format → save JSONL.
    Returns the absolute path to the saved JSONL file.
    """
    train_dir   = config.TRAIN_DIR
    data_dir    = config.DATA_DIR
    output_path = os.path.join(data_dir, "train.jsonl")

    os.makedirs(data_dir, exist_ok=True)

    print("=" * 65)
    print("  Chunking Training Data")
    print("=" * 65)

    # ── Scan for all supported file types ─────────────────────────────────
    all_files: list[str] = []
    for ext in SUPPORTED_EXTENSIONS:
        all_files.extend(
            glob.glob(os.path.join(train_dir, "**", f"*{ext}"), recursive=True)
        )
    all_files = sorted(set(all_files))  # deduplicate paths on case-insensitive FS

    if not all_files:
        print(f"\n  [WARNING] No supported files found in {os.path.abspath(train_dir)}")
        print(f"  Supported: .txt  .pdf  .docx  .xlsx  .xls  .md  .markdown")
        sys.exit(1)

    ext_counts: dict[str, int] = {}
    for f in all_files:
        ext = os.path.splitext(f)[1].lower()
        ext_counts[ext] = ext_counts.get(ext, 0) + 1

    print(f"\n  Files found: {len(all_files)}")
    for ext, count in sorted(ext_counts.items()):
        print(f"    {ext:<12} {count}")
    print(f"\n  Chunk size:    {config.CHUNK_SIZE} words")
    print(f"  Chunk overlap: {config.CHUNK_OVERLAP} words")
    print()

    # ── Extract text and chunk ─────────────────────────────────────────────
    raw_chunks: list[str] = []
    skipped = 0

    for filepath in all_files:
        rel_path  = os.path.relpath(filepath, train_dir)
        ext_label = os.path.splitext(filepath)[1].upper()[1:]
        extractor = get_extractor(filepath)
        if extractor is None:
            continue

        print(f"  [{ext_label:8}] {rel_path}", end="")
        text = extractor(filepath)
        if not text.strip():
            print(" → skipped (empty or unreadable)")
            skipped += 1
            continue

        chunks = chunk_text(text, config.CHUNK_SIZE, config.CHUNK_OVERLAP)
        print(f" → {len(chunks)} chunks")
        raw_chunks.extend(chunks)

    if not raw_chunks:
        print("\n  [ERROR] No text extracted from any files.")
        print("  Ensure files contain readable text (not scanned images).")
        sys.exit(1)

    raw_count = len(raw_chunks)

    # ── Deduplication ──────────────────────────────────────────────────────
    if config.DEDUP_ENABLED:
        print()
        print(f"  Running deduplication  (Jaccard threshold: {config.DEDUP_THRESHOLD})...")
        final_chunks, removed = deduplicate_chunks(raw_chunks, config.DEDUP_THRESHOLD)
        if removed:
            print(f"  Removed {removed} near-duplicate chunk(s)  "
                  f"({raw_count} → {len(final_chunks)})")
        else:
            print(f"  No duplicates found.")
    else:
        final_chunks = raw_chunks
        print(f"\n  Deduplication: DISABLED  (set DEDUP_ENABLED=True to enable)")

    # ── Validation ─────────────────────────────────────────────────────────
    if config.VALIDATE_CHUNKS:
        report = validate_chunks(final_chunks)
        print_validation_report(report)

    # ── Format with chat template ──────────────────────────────────────────
    print()
    print(f"  Formatting {len(final_chunks)} chunks with chat template...")
    template   = get_chat_template(config.MODEL_NAME)
    formatted_examples = []
    for chunk in final_chunks:
        formatted = format_training_example(
            template=template,
            user_text=(
                "Based on the following information, provide a helpful "
                "and detailed response:\n\n" + chunk
            ),
            assistant_text="Based on the provided information:\n\n" + chunk,
            system_text=config.SYSTEM_PROMPT,
        )
        formatted_examples.append({"text": formatted})

    # ── Save JSONL ─────────────────────────────────────────────────────────
    with open(output_path, "w", encoding="utf-8") as f:
        for item in formatted_examples:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")

    abs_output = os.path.abspath(output_path)
    print()
    print("  ┌─────────────────────────────────────────────────────────────┐")
    print(f"  │  Files processed:   {len(all_files) - skipped:<40}│")
    print(f"  │  Files skipped:     {skipped:<40}│")
    print(f"  │  Raw chunks:        {raw_count:<40}│")
    print(f"  │  After dedup:       {len(final_chunks):<40}│")
    print(f"  │  Saved to:          {abs_output:<40}│")
    print("  └─────────────────────────────────────────────────────────────┘")
    print()
    return output_path


if __name__ == "__main__":
    run_chunking()
